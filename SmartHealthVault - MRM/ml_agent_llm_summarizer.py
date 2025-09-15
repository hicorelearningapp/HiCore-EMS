import os
from ddgs import DDGS
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from langchain_openai import ChatOpenAI
from langchain.agents import initialize_agent, Tool

# -----------------------------
# Tool 1: Web Search
# -----------------------------
def search_tool(query: str):
    results = []
    with DDGS() as ddgs:
        for r in ddgs.text(query, max_results=5):
            results.append(f"{r['title']} - {r['href']}\n{r['body']}")
    return "\n\n".join(results)

# -----------------------------
# Tool 2: Summarizer
# -----------------------------
def summarize_tool(text: str):
    llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)
    prompt = f"""
    Summarize the following into a clean professional report.

    Use numbered section headings (1. Overview, 2. Advantages, etc.).  
    Write short paragraphs and bullet points if needed.  
    Format bullet points like `Key Term: Explanation` (so we can bold the Key Term).  
    Do NOT use markdown symbols (#, *, **).  
    
    Text:
    {text}
    """
    return llm.invoke(prompt).content

# -----------------------------
# Tool 3: Template Filler
# -----------------------------
def template_tool(data: dict, output_path: str = None) -> str:
    """
    Generate a Word document with the report content.
    
    Args:
        data: Dictionary containing 'topic' and 'summary'
        output_path: Optional path to save the document. If not provided, uses a generated filename.
    
    Returns:
        str: Path to the generated document
    """
    doc = Document()

    # Title
    title = doc.add_heading(f"Report on {data['topic']}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0, 51, 153)

    doc.add_paragraph("")  # spacing

    # Add Summary Sections
    summary_lines = data["summary"].splitlines()

    for line in summary_lines:
        line = line.strip()
        if not line:
            continue

        # Section headers like "1. Overview"
        if line[0].isdigit() and "." in line:
            heading = doc.add_heading(line, level=1)
            run = heading.runs[0]
            run.font.color.rgb = RGBColor(0, 51, 153)   # Blue
            run.font.bold = True
            run.font.size = Pt(14)

        # Bullet points with bolded key terms
        elif line.startswith("-"):
            para = doc.add_paragraph(style="List Bullet")
            line = line[1:].strip()
            line = line.replace("**", "")  # Remove markdown bold

            if ":" in line:
                key, rest = line.split(":", 1)
                run1 = para.add_run(key + ": ")
                run1.bold = True
                run1.font.color.rgb = RGBColor(0, 0, 128)
                run1.font.size = Pt(12)

                run2 = para.add_run(rest.strip())
                run2.font.color.rgb = RGBColor(50, 50, 50)
                run2.font.size = Pt(12)
            else:
                para.add_run(line)

        # Normal paragraph
        else:
            para = doc.add_paragraph(line)
            run = para.runs[0]
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(50, 50, 50)

    # Determine output path
    if not output_path:
        safe_topic = "".join(c if c.isalnum() or c in ' _-' else '_' for c in data['topic'])
        output_path = f"{safe_topic}_report.docx"
    
    # Ensure directory exists
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    
    # Save the document
    doc.save(output_path)
    return output_path

# -----------------------------
# Register Tools with Agent
# -----------------------------
tools = [
    Tool(name="Search", func=search_tool, description="Search the web for a topic"),
    Tool(name="Summarize", func=summarize_tool, description="Summarize content into structured format"),
    Tool(name="TemplateFiller", func=template_tool, description="Fill Word template with data"),
]

llm = ChatOpenAI(model="gpt-4o-mini", temperature=0)
agent = initialize_agent(tools, llm, agent="zero-shot-react-description", verbose=True)

# -----------------------------
# Main Execution
# -----------------------------
def generate_report(topic: str, output_path: str = None):
    raw_results = search_tool(topic)             # Step 1: Search
    summary = summarize_tool(raw_results)        # Step 2: Summarize
    result = template_tool({"topic": topic, "summary": summary}, output_path)  # Step 3: Fill Template
    return result
